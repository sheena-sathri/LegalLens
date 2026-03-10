"""Document processor — extracts text from PDF and DOCX files."""

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import os
import uuid
from typing import BinaryIO

from app.utils.text_processing import clean_text
from app.utils.chunking import chunk_by_sections
from app.db.vector_store import add_document_chunks
from app.db.metadata_store import save_document


async def process_document(file: BinaryIO, filename: str) -> dict:
    """
    Process an uploaded document:
    1. Extract text from PDF or DOCX
    2. Clean and normalize text
    3. Chunk the text respecting section boundaries
    4. Store chunks in vector DB
    5. Save metadata to SQLite

    Returns document metadata dict.
    """
    document_id = f"doc_{uuid.uuid4().hex[:12]}"
    file_type = _get_file_type(filename)
    file_bytes = file.read()

    # Extract text
    if file_type == "pdf":
        raw_text, page_count = _extract_pdf(file_bytes)
    elif file_type == "docx":
        raw_text, page_count = _extract_docx(file_bytes)
    elif file_type == "txt":
        raw_text = file_bytes.decode("utf-8", errors="ignore")
        page_count = 1
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    # Clean text
    cleaned_text = clean_text(raw_text)

    # Chunk and index
    chunks = chunk_by_sections(cleaned_text)
    num_chunks = add_document_chunks(document_id, chunks)

    # Save to SQLite
    await save_document(
        document_id=document_id,
        filename=filename,
        file_type=file_type,
        page_count=page_count,
        character_count=len(cleaned_text),
        raw_text=cleaned_text,
    )

    return {
        "document_id": document_id,
        "filename": filename,
        "file_type": file_type,
        "page_count": page_count,
        "character_count": len(cleaned_text),
        "chunks_indexed": num_chunks,
        "status": "uploaded",
    }


def _get_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".text": "txt",
    }
    return type_map.get(ext, "unknown")


def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages), len(pages) if pages else doc.page_count


def _extract_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from DOCX using python-docx."""
    import io
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    # Estimate page count (rough: ~3000 chars per page)
    page_count = max(1, len(full_text) // 3000)
    return full_text, page_count
